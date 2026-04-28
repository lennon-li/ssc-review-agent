import os
from pathlib import Path
from typing import List, Dict, Any
import json
import yaml
from .loaders import extract_all_from_folder, load_yaml, load_text

from datetime import datetime

def list_files_in_folder(folder_path: str) -> List[str]:
    """Lists all files in the specified folder."""
    path = Path(folder_path)
    if not path.exists() or not path.is_dir():
        return []
    return [f.name for f in path.iterdir() if f.is_file()]

def find_applicant_photo(folder_path: str) -> str:
    """
    Looks for the first image file in the folder.
    Returns the filename or None.
    """
    path = Path(folder_path)
    if not path.exists() or not path.is_dir():
        return None
    for f in path.iterdir():
        if f.suffix.lower() in ['.png', '.jpg', '.jpeg']:
            return f.name
    return None

def save_reviewed_evaluation(folder_name: str, evaluation_data: Dict[str, Any]) -> str:
    """
    Saves the edited evaluation data to outputs/<folder_name>/evaluation_reviewed.json.
    """
    base_dir = Path(__file__).resolve().parent.parent
    output_dir = base_dir / "outputs" / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)
    
    output_file = output_dir / "evaluation_reviewed.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(evaluation_data, f, indent=2)
    
    return str(output_file)

def generate_markdown_report(result: Dict[str, Any]) -> str:
    """Generates a Markdown string from the evaluation result."""
    md = f"# SSC Accreditation Review: {result.get('applicant_id', 'Unknown')}\n\n"
    
    recommendation = result.get('ai_recommendation', 'N/A').upper()
    md += f"## Suggested Recommendation for Human Reviewer\n"
    md += f"**{recommendation}**\n\n"
    md += "*Note: This is an AI-generated suggestion based on the initial analysis. The final decision remains with the human reviewer.*\n\n"

    flags = result.get('ai_flags', [])
    if flags:
        md += "## [!] AI Flags & Assumptions\n"
        for flag in flags:
            md += f"- **{flag['topic']}**: {flag['reason']} (Suggestion: {flag['suggestion']})\n"
        md += "\n"

    md += f"## Executive Summary\n{result.get('overall_summary', 'No summary provided.')}\n\n"
    
    md += "## Course Checklist\n"
    checklist = result.get('course_checklist', [])
    if checklist:
        md += "| Module | Course Code | Status |\n|---|---|---|\n"
        for item in checklist:
            status = "✅ Satisfied" if item.get('is_satisfied') else "❌ Missing/Fail"
            md += f"| {item.get('module')} | {item.get('course_code')} | {status} |\n"
    else:
        md += "No course data provided.\n"
    md += "\n"
    
    md += "## Detailed Criterion Assessment\n"
    for crit in result.get('criteria', []):
        status = "⚠️ Needs Attention" if crit.get('needs_human_attention') else "✅ Satisfied"
        md += f"### {crit.get('criterion_name')} ({status})\n"
        md += f"{crit.get('supporting_evidence', 'No evidence provided.')}\n\n"
        
    return md

def generate_latex_report(result: Dict[str, Any]) -> str:
    """Generates a basic LaTeX string from the evaluation result."""
    tex = r"\documentclass{article}" + "\n"
    tex += r"\usepackage[utf8]{inputenc}" + "\n"
    tex += r"\usepackage{geometry}" + "\n"
    tex += r"\geometry{margin=1in}" + "\n"
    tex += r"\title{SSC Accreditation Review: " + result.get('applicant_id', 'Unknown') + "}\n"
    tex += r"\author{SSC Review Agent}" + "\n"
    tex += r"\date{\today}" + "\n"
    tex += r"\begin{document}" + "\n"
    tex += r"\maketitle" + "\n\n"
    
    recommendation = result.get('ai_recommendation', 'N/A').upper()
    tex += r"\section*{Suggested Recommendation for Human Reviewer}" + "\n"
    tex += r"\textbf{" + recommendation + "}\n\n"
    tex += r"\textit{Note: This is an AI-generated suggestion based on the initial analysis. The final decision remains with the human reviewer.}\n\n"

    flags = result.get('ai_flags', [])
    if flags:
        tex += r"\section*{AI Flags \& Assumptions}" + "\n"
        tex += r"\begin{itemize}" + "\n"
        for flag in flags:
            topic = flag['topic'].replace("_", r"\_").replace("&", r"\&")
            reason = flag['reason'].replace("_", r"\_").replace("&", r"\&")
            suggestion = flag['suggestion'].replace("_", r"\_").replace("&", r"\&")
            tex += f"  \\item \\textbf{{{topic}}}: {reason} -- \\textit{{Suggestion: {suggestion}}}\n"
        tex += r"\end{itemize}" + "\n\n"

    tex += r"\section*{Executive Summary}" + "\n"
    tex += result.get('overall_summary', 'No summary provided.').replace("_", r"\_").replace("&", r"\&") + "\n\n"
    
    tex += r"\section*{Course Checklist}" + "\n"
    checklist = result.get('course_checklist', [])
    if checklist:
        tex += r"\begin{itemize}" + "\n"
        for item in checklist:
            status = "Satisfied" if item.get('is_satisfied') else "Missing/Fail"
            tex += f"  \\item {item.get('module')}: {item.get('course_code')} -- {status}\n"
        tex += r"\end{itemize}" + "\n"
    else:
        tex += "No course data provided.\n"
    
    tex += r"\section*{Detailed Criterion Assessment}" + "\n"
    for crit in result.get('criteria', []):
        tex += r"\subsection*{" + crit.get('criterion_name') + "}\n"
        tex += crit.get('supporting_evidence', 'No evidence provided.').replace("_", r"\_").replace("&", r"\&") + "\n\n"
        
    tex += r"\end{document}"
    return tex

def generate_docx_report(result: Dict[str, Any], output_path: str):
    """Generates a DOCX file from the evaluation result."""
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    doc.add_heading(f"SSC Accreditation Review: {result.get('applicant_id', 'Unknown')}", 0)
    
    doc.add_heading("Suggested Recommendation for Human Reviewer", level=1)
    recommendation = result.get('ai_recommendation', 'N/A').upper()
    p = doc.add_paragraph()
    run = p.add_run(recommendation)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("Note: This is an AI-generated suggestion based on the initial analysis. The final decision remains with the human reviewer.")

    flags = result.get('ai_flags', [])
    if flags:
        doc.add_heading("AI Flags & Assumptions", level=1)
        for flag in flags:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{flag['topic']}: ")
            run.bold = True
            p.add_run(f"{flag['reason']} (Suggestion: {flag['suggestion']})")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(result.get('overall_summary', 'No summary provided.'))
    
    doc.add_heading("Course Checklist", level=1)
    checklist = result.get('course_checklist', [])
    if checklist:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Module'
        hdr_cells[1].text = 'Course Code'
        hdr_cells[2].text = 'Status'
        for item in checklist:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('module'))
            row_cells[1].text = str(item.get('course_code'))
            row_cells[2].text = "Satisfied" if item.get('is_satisfied') else "Missing/Fail"
    else:
        doc.add_paragraph("No course data provided.")
        
    doc.add_heading("Detailed Criterion Assessment", level=1)
    for crit in result.get('criteria', []):
        status = " (Needs Attention)" if crit.get('needs_human_attention') else ""
        doc.add_heading(f"{crit.get('criterion_name')}{status}", level=2)
        doc.add_paragraph(crit.get('supporting_evidence', 'No evidence provided.'))
        
    doc.save(output_path)
    return output_path

def run_folder_evaluation(folder_path: str, evaluator_type: str = "vertex") -> Dict[str, Any]:
    """
    High-level entry point for Shiny. 
    Extracts text, loads rubric/instructions, and runs evaluation.
    Includes metadata in the result.
    """
    from .evaluate import get_evaluator
    
    # 1. Extract
    app_text = extract_all_from_folder(folder_path)
    files_processed = list_files_in_folder(folder_path)
    
    # 2. Load context
    base_dir = Path(__file__).resolve().parent.parent
    rubric_path = base_dir / "criteria" / "rubric.yml"
    rubric = load_yaml(str(rubric_path))
    
    # Only load instructions if NOT using vertex (to avoid prompt stuffing)
    instructions = ""
    instructions_path = None
    if evaluator_type != "vertex":
        instructions_path = base_dir / "skills" / "ssc-evaluator" / "SKILL.md"
        if not instructions_path.exists():
            instructions_path = base_dir / "instructions" / "reviewer_instructions.md"
        
        if instructions_path.exists():
            instructions = load_text(str(instructions_path))
    
    # 3. Evaluate
    evaluator = get_evaluator(evaluator_type)
    result = evaluator.evaluate(app_text, rubric, instructions)
    
    # 4. Add Metadata
    result["metadata"] = {
        "timestamp": datetime.now().isoformat(),
        "evaluator_type": evaluator_type,
        "model_id": getattr(evaluator, "model_id", "unknown"),
        "applicant_folder_path": folder_path,
        "files_processed": files_processed,
        "rubric_path": str(rubric_path),
        "instructions_path": str(instructions_path) if instructions_path else "Retrieved from Data Store"
    }
    
    # 5. Save
    folder_name = Path(folder_path).name
    output_dir = base_dir / "outputs" / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)
    
    output_file = output_dir / "evaluation_result.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2)
        
    with open(output_dir / "extracted_content.txt", "w", encoding="utf-8") as f:
        f.write(app_text)
        
    return result
