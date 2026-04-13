import docx
import sys

def convert_docx_to_md(docx_path):
    doc = docx.Document(docx_path)
    md_lines = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = para.style.name.split()[-1]
            if level.isdigit():
                md_lines.append('#' * int(level) + ' ' + para.text)
            else:
                md_lines.append('# ' + para.text)
        elif para.text.strip():
            md_lines.append(para.text)
            
    # Process tables (crucial for review forms)
    for table in doc.tables:
        md_lines.append("\n--- Table ---")
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_lines.append(" | ".join(row_text))
            
    return '\n\n'.join(md_lines)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    print(convert_docx_to_md(docx_path))
